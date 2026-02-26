#!/usr/bin/env python3
"""
embed_images.py — Post-process a rendered DOCX to replace [[SECTION:*]] markers
with grouped, embedded photo sections using python-docx.

Usage:
  python3 scripts/embed_images.py \
    --docx <rendered.docx> \
    --photos <classified_photos.json> \
    --output <output.docx> \
    [--max-width-cm 15]

The photos JSON must already be classified (run classify_photos.py first, or pass
the combined classify+embed pipeline).
"""

import argparse
import json
import os
import shutil
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from PIL import Image as PilImage
    HAS_PILLOW = True
except ImportError:
    HAS_PILLOW = False

# ─────────────────────── Constants ────────────────────────────────────────────

SECTION_HEADINGS = {
    'site':      'Site Photographs',
    'guideline': 'Guideline Value Screenshot(s)',
    'map':       'Route Map / Google Maps',
    'dishank':   'GPS / Dishank Location',
}

SECTION_ORDER = ['site', 'guideline', 'map', 'dishank']
MARKER_PREFIX = '[[SECTION:'
COVER_MARKER = '[[COVER_IMAGE]]'

# ─────────────────────── Helpers ─────────────────────────────────────────────

def find_marker_paragraphs(doc: 'Document') -> dict[str, object]:
    """Return {marker_text: paragraph} for all [[...]] markers, searching
    entire document body (main body + tables)."""
    markers = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith('[[') and text.endswith(']]'):
            markers[text] = para
    # Also scan table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text.startswith('[[') and text.endswith(']]'):
                        markers[text] = para
    return markers


def insert_section_after(doc: 'Document', marker_para, category: str,
                         photos: list[dict], max_width_cm: float) -> None:
    """Insert heading + images directly after (and including replacement of) marker."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    # Get the XML element for the marker paragraph
    marker_elem = marker_para._element
    parent = marker_elem.getparent()
    idx = list(parent).index(marker_elem)

    # Elements to insert (in reverse order since we insert at same idx)
    new_elems = []

    # ── Heading paragraph ──
    heading_para = doc.add_paragraph()
    heading_para.clear()
    run = heading_para.add_run(SECTION_HEADINGS.get(category, category.title()))
    run.bold = True
    run.font.size = Pt(12)
    heading_para.paragraph_format.space_before = Pt(12)
    heading_para.paragraph_format.space_after = Pt(6)
    new_elems.append(copy.deepcopy(heading_para._element))

    # ── Image paragraphs ──
    for photo in photos:
        photo_path = photo.get('path', '')
        if not photo_path or not os.path.exists(photo_path):
            # Still add caption even if file missing
            cap_para = doc.add_paragraph()
            cap_para.add_run(f"[Image not found: {photo.get('filename', 'unknown')}]").italic = True
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            new_elems.append(copy.deepcopy(cap_para._element))
            continue

        # Optionally resize to fit max width
        resolved_path = photo_path
        if HAS_PILLOW:
            resolved_path = maybe_resize(photo_path, max_width_cm)

        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run_img = img_para.add_run()
            run_img.add_picture(resolved_path, width=Cm(max_width_cm))
        except Exception as e:
            img_para.clear()
            img_para.add_run(f"[Image error: {photo.get('filename', 'unknown')}: {e}]").italic = True
        new_elems.append(copy.deepcopy(img_para._element))

        # Caption
        caption = photo.get('caption', '')
        if caption:
            cap_para = doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_para.add_run(caption)
            cap_run.italic = True
            cap_run.font.size = Pt(9)
            cap_para.paragraph_format.space_after = Pt(8)
            new_elems.append(copy.deepcopy(cap_para._element))

    # Insert all new elements after marker position
    for i, elem in enumerate(new_elems):
        parent.insert(idx + 1 + i, elem)

    # Remove the marker paragraph
    parent.remove(marker_elem)

    # Clean up temporary paragraphs added to doc (they were deep-copied, so remove from doc body)
    for extra_para in doc.paragraphs[-len(new_elems):]:
        p = extra_para._element
        p.getparent().remove(p)


def maybe_resize(path: str, max_width_cm: float) -> str:
    """Return path; if image wider than max, return a resized temp copy."""
    try:
        with PilImage.open(path) as img:
            w, h = img.size
            dpi = img.info.get('dpi', (96, 96))[0] or 96
            w_cm = (w / dpi) * 2.54
            if w_cm <= max_width_cm:
                return path
            # Scale down
            scale = max_width_cm / w_cm
            new_w = int(w * scale)
            new_h = int(h * scale)
            resized = img.resize((new_w, new_h), PilImage.LANCZOS)
            tmp = f"/tmp/zenops_resize_{Path(path).stem}.jpg"
            resized.save(tmp, 'JPEG', quality=90)
            return tmp
    except Exception:
        return path


# ─────────────────────── Main ─────────────────────────────────────────────────

def embed(docx_path: str, photos: list[dict], output_path: str,
          max_width_cm: float = 14.5) -> dict:
    """Core embed function. Returns stats dict."""
    if not HAS_DOCX:
        raise RuntimeError("python-docx not installed. Run: pip3 install python-docx")

    doc = Document(docx_path)
    markers = find_marker_paragraphs(doc)

    # Group photos by category
    groups: dict[str, list[dict]] = {cat: [] for cat in SECTION_ORDER}
    for photo in photos:
        cat = photo.get('category', 'site')
        if cat in groups:
            groups[cat].append(photo)

    stats = {'sections_inserted': 0, 'sections_skipped': 0, 'photos_embedded': 0}

    for cat in SECTION_ORDER:
        marker_key = f'[[SECTION:{cat.upper()}]]'
        # Also check variations
        candidates = [marker_key,
                      f'[[SECTION:{cat.upper().replace("_", "-")}]]',
                      f'[[SECTION:PHOTOS_{cat.upper()}]]']

        para = None
        for cand in candidates:
            if cand in markers:
                para = markers[cand]
                break

        cat_photos = groups.get(cat, [])

        if para is None:
            # No marker for this section (template doesn't have it)
            if cat_photos:
                print(f'  [warn] No marker found for [{cat}] — {len(cat_photos)} photos will be appended at end')
            continue

        if not cat_photos:
            # Omit empty section — remove marker paragraph
            marker_elem = para._element
            marker_elem.getparent().remove(marker_elem)
            stats['sections_skipped'] += 1
            continue

        insert_section_after(doc, para, cat, cat_photos, max_width_cm)
        stats['sections_inserted'] += 1
        stats['photos_embedded'] += len(cat_photos)
        print(f'  [{cat}] {len(cat_photos)} photo(s) embedded')

    # Ensure output directory exists
    os.makedirs(Path(output_path).parent, exist_ok=True)
    if output_path == docx_path:
        tmp = docx_path + '.tmp'
        doc.save(tmp)
        shutil.move(tmp, output_path)
    else:
        doc.save(output_path)

    return stats


def main():
    parser = argparse.ArgumentParser(description='ZenOps DOCX image embedder')
    parser.add_argument('--docx', required=True)
    parser.add_argument('--photos', required=True, help='Classified photos JSON')
    parser.add_argument('--output', required=True)
    parser.add_argument('--max-width-cm', type=float, default=14.5)
    args = parser.parse_args()

    with open(args.photos) as f:
        photos = json.load(f)

    stats = embed(args.docx, photos, args.output, args.max_width_cm)
    print(f'Done. {stats["sections_inserted"]} section(s) inserted, '
          f'{stats["sections_skipped"]} empty section(s) removed, '
          f'{stats["photos_embedded"]} photo(s) embedded.')


if __name__ == '__main__':
    main()
